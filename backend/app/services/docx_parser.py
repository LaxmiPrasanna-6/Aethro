"""
Word (.docx) document parser for bulk room data extraction.

Handles ALL common Word document layouts:
  1. Plain paragraphs with blank-line separators  (most common)
  2. Content inside Word tables                   (very common with formatted docs)
  3. Compact paragraphs with NO blank lines       (fallback: split on "Room ID:" keyword)
  4. Non-breaking spaces (\xa0) used as blank     (Word default behaviour)
  5. Mixed bold/italic runs — only text is read   (formatting stripped automatically)

Accepted document format:

    Block: A
    Room ID: A-101
    Type: Classroom
    Capacity: 30
    Facilities: projector, ac, whiteboard

    Block: A
    Room ID: A-102
    Type: Lab
    Capacity: 25
    Facilities: computers, ac
"""
import io
import re
from typing import Optional
from docx import Document


# ─── Type normalisation ────────────────────────────────────────────────────────

VALID_TYPES = {"classroom", "lab", "seminar_hall"}

TYPE_ALIASES: dict[str, str] = {
    "classroom":       "classroom",
    "class room":      "classroom",
    "class":           "classroom",
    "lecture hall":    "classroom",
    "lecture room":    "classroom",
    "lab":             "lab",
    "laboratory":      "lab",
    "computer lab":    "lab",
    "it lab":          "lab",
    "seminar hall":    "seminar_hall",
    "seminar_hall":    "seminar_hall",
    "seminar":         "seminar_hall",
    "auditorium":      "seminar_hall",
    "conference":      "seminar_hall",
    "conference hall": "seminar_hall",
    "conference room": "seminar_hall",
    "hall":            "seminar_hall",
}

# ─── Field aliases (case-insensitive) ─────────────────────────────────────────

FIELD_ALIASES: dict[str, str] = {
    "block":         "block",
    "block no":      "block",
    "block number":  "block",
    "room id":       "room_id",
    "roomid":        "room_id",
    "room no":       "room_id",
    "room number":   "room_id",
    "room":          "room_id",
    "id":            "room_id",
    "type":          "type",
    "room type":     "type",
    "capacity":      "capacity",
    "cap":           "capacity",
    "seats":         "capacity",
    "max capacity":  "capacity",
    "strength":      "capacity",
    "facilities":    "facilities",
    "facility":      "facilities",
    "amenities":     "facilities",
    "features":      "facilities",
    "equipment":     "facilities",
    "floor":         "floor",
    "floor number":  "floor",
    "floor no":      "floor",
    "beds":          "beds",
    "number of beds": "beds",
    "num beds":      "beds",
}

# Hostel/Lodge specific field aliases
HOSTEL_LODGE_FIELD_ALIASES: dict[str, str] = {
    "room id":       "room_id",
    "roomid":        "room_id",
    "room no":       "room_id",
    "room number":   "room_id",
    "room":          "room_id",
    "id":            "room_id",
    "floor":         "floor",
    "floor number":  "floor",
    "floor no":      "floor",
    "beds":          "beds",
    "number of beds": "beds",
    "num beds":      "beds",
    "capacity":      "beds",
    "features":      "features",
    "amenities":     "features",
}

REQUIRED_HOSTEL_LODGE_FIELDS = {"room_id", "floor", "beds"}

REQUIRED_FIELDS = {"room_id", "type", "capacity"}

STUDENT_FIELD_ALIASES: dict[str, str] = {
    "student id":    "student_id",
    "student_id":    "student_id",
    "studentid":     "student_id",
    "roll no":       "roll_number",
    "roll number":   "roll_number",
    "roll":          "roll_number",
    "name":          "name",
    "student name":  "name",
    "department":    "department",
    "dept":          "department",
    "semester":      "semester",
    "sem":           "semester",
    "subject":       "subjects",
    "subjects":      "subjects",
    "email":         "email",
}

REQUIRED_STUDENT_FIELDS = {"student_id", "name", "roll_number", "department", "semester", "subjects"}

# Lines that definitely start a new room block (used as fallback splitter)
BLOCK_START_PATTERN = re.compile(
    r"^\s*(room\s*id|room\s*no|block)\s*:", re.IGNORECASE
)


# ─── Low-level helpers ─────────────────────────────────────────────────────────

def _is_blank(line: str) -> bool:
    """
    True when a line carries no meaningful text.
    Covers: empty string, only spaces/tabs, only non-breaking spaces (\xa0),
    only Unicode whitespace, only bullet/dash characters.
    """
    cleaned = (
        line
        .replace("\xa0", "")   # non-breaking space (Word default)
        .replace("\u200b", "") # zero-width space
        .replace("\ufeff", "") # BOM
        .strip("•·–—- \t\r\n")
    )
    return cleaned == ""


def _normalise_type(raw: str) -> Optional[str]:
    return TYPE_ALIASES.get(raw.strip().lower())


def _normalise_field_key(raw: str) -> Optional[str]:
    return FIELD_ALIASES.get(raw.strip().lower())


def _parse_facilities(raw: str) -> list[str]:
    """Split comma/semicolon/pipe-separated facility list, lowercase each entry."""
    parts = re.split(r"[,;|]", raw)
    return [p.strip().lower() for p in parts if p.strip()]

def _parse_subjects(raw: str) -> list[str]:
    """Split comma/semicolon/pipe-separated subject list."""
    parts = re.split(r"[,;|]", raw)
    return [p.strip() for p in parts if p.strip()]


def _normalise_student_field_key(raw: str) -> Optional[str]:
    return STUDENT_FIELD_ALIASES.get(raw.strip().lower())


def _normalise_hostel_lodge_field_key(raw: str) -> Optional[str]:
    return HOSTEL_LODGE_FIELD_ALIASES.get(raw.strip().lower())

# ─── Extraction from Word document ────────────────────────────────────────────

def _para_text(para) -> str:
    """Return the full text of a paragraph by joining all its runs."""
    return "".join(run.text for run in para.runs) or para.text


def _unique_cells(row) -> list[str]:
    """
    Return deduplicated cell texts for a row, ignoring merged-cell duplicates.
    python-docx repeats the same cell object for each merged column span;
    we use the underlying XML element identity to collapse those duplicates.
    """
    seen = set()
    texts = []
    for cell in row.cells:
        cid = id(cell._tc)
        if cid in seen:
            continue
        seen.add(cid)
        t = cell.text.strip()
        if t:
            texts.append(t)
    return texts


def extract_all_lines(file_bytes: bytes) -> list[str]:
    """
    Extract every line of text from a .docx file.

    Strategy (in order):
      1. Regular body paragraphs.
      2. Cells inside every table.
         • 2-column table  (Field | Value per row) → "Field: Value" lines grouped
           per room with blank separators.
         • Header + data table (spreadsheet style, ≥2 cols, first row = field names)
           → each data row becomes "Header: Value" lines separated by blanks.
         • Single-column multi-line cell → splitlines().
         • Other → raw cell texts (legacy behaviour).
    """
    doc   = Document(io.BytesIO(file_bytes))
    lines: list[str] = []

    # ── Regular paragraphs ────────────────────────────────────────────────────
    for para in doc.paragraphs:
        lines.append(_para_text(para))

    # ── Tables ────────────────────────────────────────────────────────────────
    for table in doc.tables:
        all_rows = table.rows

        # --- Detect header+data layout (spreadsheet-style) -------------------
        # First row cells must all resolve to a known FIELD_ALIASES key.
        header_cells = _unique_cells(all_rows[0]) if all_rows else []
        header_fields = [_normalise_field_key(h) for h in header_cells]
        is_header_table = (
            len(header_cells) >= 2
            and all(f is not None for f in header_fields)
        )

        if is_header_table:
            for row in all_rows[1:]:            # skip the header row itself
                data_cells = _unique_cells(row)
                if not data_cells:
                    continue
                lines.append("")
                for hdr, val in zip(header_cells, data_cells):
                    if val:
                        lines.append(f"{hdr}: {val}")
                lines.append("")
            continue                            # done with this table

        # --- Row-by-row processing for non-header tables ----------------------
        for row in all_rows:
            cell_texts = _unique_cells(row)

            if not cell_texts:
                continue

            # Single cell with internal newlines → expand into lines
            if len(cell_texts) == 1 and "\n" in cell_texts[0]:
                lines.append("")
                lines.extend(cell_texts[0].splitlines())
                lines.append("")

            # 2-column Field | Value row → join as "Field: Value"
            elif len(cell_texts) == 2:
                key, val = cell_texts
                # Only join if the key looks like a field label (no colon yet)
                if ":" not in key:
                    lines.append(f"{key}: {val}")
                else:
                    # Already "Key: something" — keep as-is
                    lines.append(key)
                    lines.append(val)

            # All other multi-cell rows → raw cell texts (legacy)
            else:
                lines.append("")
                lines.extend(cell_texts)
                lines.append("")

    return lines


# ─── Block splitting ──────────────────────────────────────────────────────────

def _split_by_blank_lines(lines: list[str]) -> list[list[str]]:
    """Standard split: blank line = block boundary."""
    blocks:  list[list[str]] = []
    current: list[str] = []
    for line in lines:
        if _is_blank(line):
            if current:
                blocks.append(current)
                current = []
        else:
            current.append(line.strip())
    if current:
        blocks.append(current)
    return blocks


def _split_by_room_keyword(lines: list[str]) -> list[list[str]]:
    """
    Fallback split: start a new block each time 'Room ID:' or 'Block:' appears.
    Used when the document has no blank-line separators at all.
    """
    blocks:  list[list[str]] = []
    current: list[str] = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            # Include blank lines in current block (they may be meaningful separators)
            if current:
                current.append("")
            continue
        # Check if this line starts a new block (Room ID: or Block:)
        if BLOCK_START_PATTERN.match(stripped):
            if current:
                # Append previous block
                blocks.append(current)
                current = []
            # Start new block with this line
            current.append(stripped)
        else:
            # Continue current block
            current.append(stripped)
    
    if current:
        blocks.append(current)
    return blocks


def split_into_blocks(lines: list[str]) -> list[list[str]]:
    """
    Try blank-line splitting first.
    If that yields ≤1 block, fall back to keyword-boundary splitting.
    """
    # First, filter out heading-like lines (all caps or ends with colon but no value)
    filtered_lines = []
    for line in lines:
        stripped = line.strip()
        # Skip obvious headings (all caps, or empty)
        if not stripped or (stripped.isupper() and len(stripped) > 3):
            continue
        filtered_lines.append(line)
    
    blocks = _split_by_blank_lines(filtered_lines)
    if len(blocks) <= 1:
        blocks = _split_by_room_keyword(filtered_lines)
    return blocks


# ─── Block parsing & validation ───────────────────────────────────────────────

class ParseError(Exception):
    """Raised when a room block has unrecoverable data issues."""


def _parse_room_block(lines: list[str]) -> dict:
    """
    Parse a list of non-blank lines into a raw key→value dict.
    Each line is expected to be in the form 'Key: Value'.
    Lines without a colon are silently ignored (headings, decorations, etc.).
    """
    raw: dict = {}
    for line in lines:
        if ":" not in line:
            continue
        key_part, _, val_part = line.partition(":")
        field = _normalise_field_key(key_part)
        if field:
            raw[field] = val_part.strip()
    return raw


def validate_and_build_room(raw: dict, block_index: int) -> dict:
    """
    Validate the raw key-value dict and return a clean room document.
    Raises ParseError with a human-readable message on failure.
    """
    missing = REQUIRED_FIELDS - raw.keys()
    if missing:
        raise ParseError(
            f"Entry {block_index + 1}: missing required field(s): "
            f"{', '.join(sorted(missing))}"
        )

    # Capacity: must be a positive integer
    try:
        capacity = int(str(raw["capacity"]).strip())
        if capacity <= 0:
            raise ValueError
    except ValueError:
        raise ParseError(
            f"Entry {block_index + 1}: 'Capacity' must be a positive integer, "
            f"got '{raw['capacity']}'"
        )

    # Type: must be a recognised alias
    normalised_type = _normalise_type(raw["type"])
    if normalised_type is None:
        raise ParseError(
            f"Entry {block_index + 1}: unrecognised Type '{raw['type']}'. "
            f"Allowed values: Classroom, Lab, Seminar Hall  "
            f"(and common aliases like Laboratory, Auditorium, etc.)"
        )

    return {
        "block":         raw.get("block", "").strip() or None,
        "room_id":       raw["room_id"].strip(),
        "resource_type": normalised_type,
        "capacity":      capacity,
        "facilities":    _parse_facilities(raw.get("facilities", "")),
    }


# ─── Public API ───────────────────────────────────────────────────────────────

def parse_docx_rooms(file_bytes: bytes) -> dict:
    """
    Full pipeline: bytes → lines → blocks → validated room dicts.

    Returns:
        {
            "rooms":         list of valid room dicts ready for DB insertion,
            "errors":        list of human-readable error strings,
            "total_blocks":  int — number of room-like blocks found,
            "raw_preview":   first 20 extracted lines (for debugging),
        }
    """
    lines       = extract_all_lines(file_bytes)
    raw_preview = [l for l in lines if not _is_blank(l)][:20]

    blocks = split_into_blocks(lines)

    if not blocks:
        return {
            "rooms":        [],
            "errors":       [
                "No room entries found. "
                "Make sure the document uses 'Key: Value' lines "
                "(e.g. 'Room ID: A-101', 'Type: Classroom')."
            ],
            "total_blocks": 0,
            "raw_preview":  raw_preview,
        }

    rooms:  list[dict] = []
    errors: list[str]  = []

    for idx, block in enumerate(blocks):
        # Skip blocks with no colon at all (headings, decorators, etc.)
        if not any(":" in line for line in block):
            continue

        raw = _parse_room_block(block)

        # Skip blocks where NO recognised field key was found
        # (e.g. an instruction paragraph that happens to contain colons)
        if not raw:
            continue

        try:
            room = validate_and_build_room(raw, idx)
            rooms.append(room)
        except ParseError as exc:
            errors.append(str(exc))
        except Exception as exc:
            errors.append(f"Entry {idx + 1}: unexpected error — {exc}")

    return {
        "rooms":        rooms,
        "errors":       errors,
        "total_blocks": len(blocks),
        "raw_preview":  raw_preview,
    }


def _parse_student_block(lines: list[str]) -> dict:
    raw: dict = {}
    for line in lines:
        if ":" not in line:
            continue
        key_part, _, val_part = line.partition(":")
        field = _normalise_student_field_key(key_part)
        if field:
            raw[field] = val_part.strip()
    return raw


def validate_and_build_student(raw: dict, block_index: int) -> dict:
    missing = REQUIRED_STUDENT_FIELDS - raw.keys()
    if missing:
        raise ParseError(
            f"Entry {block_index + 1}: missing required field(s): {', '.join(sorted(missing))}"
        )

    try:
        semester = int(str(raw["semester"]).strip())
        if semester <= 0:
            raise ValueError
    except ValueError:
        raise ParseError(
            f"Entry {block_index + 1}: 'Semester' must be a positive integer, got '{raw['semester']}'"
        )

    subjects = _parse_subjects(raw["subjects"])
    if not subjects:
        raise ParseError(
            f"Entry {block_index + 1}: 'Subjects' must contain at least one subject"
        )

    return {
        "student_id": raw["student_id"].strip(),
        "name": raw["name"].strip(),
        "roll_number": raw["roll_number"].strip(),
        "department": raw["department"].strip(),
        "semester": semester,
        "subjects": subjects,
        "email": raw.get("email", "").strip() or None,
    }


def parse_docx_students(file_bytes: bytes) -> dict:
    lines = extract_all_lines(file_bytes)
    raw_preview = [l for l in lines if not _is_blank(l)][:20]
    blocks = split_into_blocks(lines)

    if not blocks:
        return {
            "students": [],
            "errors": [
                "No student entries found. Make sure the document uses 'Key: Value' lines"
                " (e.g. 'Roll No: 123', 'Department: CSE')."
            ],
            "total_blocks": 0,
            "raw_preview": raw_preview,
        }

    students: list[dict] = []
    errors: list[str] = []

    for idx, block in enumerate(blocks):
        if not any(":" in line for line in block):
            continue

        raw = _parse_student_block(block)
        if not raw:
            continue

        try:
            student = validate_and_build_student(raw, idx)
            students.append(student)
        except ParseError as exc:
            errors.append(str(exc))
        except Exception as exc:
            errors.append(f"Entry {idx + 1}: unexpected error — {exc}")

    return {
        "students": students,
        "errors": errors,
        "total_blocks": len(blocks),
        "raw_preview": raw_preview,
    }


def _parse_hostel_lodge_room_block(lines: list[str]) -> dict:
    """Parse a list of lines into a raw key→value dict for hostel/lodge rooms."""
    raw: dict = {}
    for line in lines:
        if ":" not in line:
            continue
        key_part, _, val_part = line.partition(":")
        field = _normalise_hostel_lodge_field_key(key_part)
        if field:
            # Clean up the value: remove extra whitespace and strip
            val = val_part.strip()
            # For features, split by comma and rejoin (handle multiple formats)
            if field == "features" and val:
                val = ", ".join(v.strip() for v in val.split(","))
            raw[field] = val
    return raw


def validate_and_build_hostel_lodge_room(raw: dict, block_index: int) -> dict:
    """Validate hostel/lodge room data."""
    missing = REQUIRED_HOSTEL_LODGE_FIELDS - raw.keys()
    if missing:
        raise ParseError(
            f"Entry {block_index + 1}: missing required field(s): "
            f"{', '.join(sorted(missing))}"
        )

    # Floor: must be a non-negative integer
    try:
        floor = int(str(raw["floor"]).strip())
        if floor < 0:
            raise ValueError
    except ValueError:
        raise ParseError(
            f"Entry {block_index + 1}: 'Floor' must be a non-negative integer, "
            f"got '{raw['floor']}'"
        )

    # Beds: must be a positive integer
    try:
        beds = int(str(raw["beds"]).strip())
        if beds <= 0:
            raise ValueError
    except ValueError:
        raise ParseError(
            f"Entry {block_index + 1}: 'Beds' must be a positive integer, "
            f"got '{raw['beds']}'"
        )

    return {
        "room_id":       raw["room_id"].strip(),
        "floor":         floor,
        "capacity":      beds,  # beds = capacity
        "features":      [f.strip() for f in raw.get("features", "").split(",") if f.strip()],
    }


def parse_docx_hostel_lodge_rooms(file_bytes: bytes) -> dict:
    """
    Parse hostel/lodge rooms from Word document.
    Expected format:
        Room ID: 101
        Floor: 1
        Beds: 2
        Features: AC, WiFi

        Room ID: 102
        Floor: 1
        Beds: 3
        Features: AC
    """
    lines       = extract_all_lines(file_bytes)
    raw_preview = [l for l in lines if not _is_blank(l)][:30]

    blocks = split_into_blocks(lines)

    rooms:  list[dict] = []
    errors: list[str]  = []
    
    # Track which blocks had colons but were invalid
    blocks_with_colons = 0
    blocks_with_valid_fields = 0

    for idx, block in enumerate(blocks):
        # Skip blocks with no colon (likely headers or decorations)
        has_colon = any(":" in line for line in block)
        if not has_colon:
            continue
        
        blocks_with_colons += 1
        raw = _parse_hostel_lodge_room_block(block)
        
        if not raw:
            # Block had colons but no recognised fields
            continue
        
        blocks_with_valid_fields += 1
        try:
            room = validate_and_build_hostel_lodge_room(raw, idx)
            rooms.append(room)
        except ParseError as exc:
            errors.append(str(exc))
        except Exception as exc:
            errors.append(f"Entry {idx + 1}: unexpected error — {exc}")

    # If no rooms were found
    if not rooms:
        error_msg = "No valid room entries found in the document."
        if blocks_with_colons == 0:
            error_msg += f"\nFound 0 blocks with colons. Document contains: {', '.join(raw_preview[:3]) if raw_preview else 'empty content'}"
        elif blocks_with_valid_fields == 0:
            error_msg += f"\nFound {blocks_with_colons} block(s) with colons but none with recognised room fields (Room ID, Floor, Beds)."
            error_msg += f"\nDocument content: {', '.join(raw_preview[:5])}"
        else:
            error_msg += f"\nAll {blocks_with_valid_fields} entries had validation errors."
        
        if not errors:
            errors.append(error_msg)

    return {
        "rooms":        rooms,
        "errors":       errors,
        "total_blocks": len(blocks),
        "raw_preview":  raw_preview,
    }


# ─── Template generator ───────────────────────────────────────────────────────

def generate_template_docx() -> bytes:
    """
    Generate a ready-to-use .docx template with sample room entries
    and a format guide at the top.
    """
    from docx import Document as _Doc
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = _Doc()

    # Title
    title = doc.add_heading("Room Upload Template", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Instructions
    doc.add_paragraph(
        "Instructions:\n"
        "• Each room entry must start with 'Room ID:' and include 'Type:' and 'Capacity:'.\n"
        "• Separate each room entry with ONE blank line.\n"
        "• Valid Type values: Classroom, Lab, Seminar Hall\n"
        "• Facilities are comma-separated (optional).\n"
        "• Do NOT use tables — keep plain Key: Value format."
    )
    doc.add_paragraph("")   # blank separator

    # Sample entries
    entries = [
        ("A", "A-101", "Classroom",   "30",  "projector, ac, whiteboard"),
        ("A", "A-102", "Lab",         "25",  "computers, ac"),
        ("A", "A-103", "Seminar Hall","100", "projector, audio_system, ac"),
        ("B", "B-201", "Classroom",   "40",  "projector, whiteboard"),
        ("B", "B-202", "Lab",         "20",  "computers"),
        ("B", "B-301", "Seminar Hall","80",  "projector, audio_system"),
    ]

    for block, room_id, rtype, capacity, facilities in entries:
        doc.add_paragraph(f"Block: {block}")
        doc.add_paragraph(f"Room ID: {room_id}")
        doc.add_paragraph(f"Type: {rtype}")
        doc.add_paragraph(f"Capacity: {capacity}")
        doc.add_paragraph(f"Facilities: {facilities}")
        doc.add_paragraph("")   # blank line between entries

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_student_template_docx() -> bytes:
    """
    Generate a ready-to-use .docx template for bulk student registration.
    """
    from docx import Document as _Doc
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = _Doc()
    title = doc.add_heading("Student Upload Template", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "Instructions:\n"
        "• Each student entry must include Student ID, Name, Roll No, Department, Semester, and Subjects.\n"
        "• Separate each student entry with ONE blank line.\n"
        "• Subjects should be comma-separated.\n"
        "• Email is optional.\n"
        "• Do NOT use tables — keep plain Key: Value format."
    )
    doc.add_paragraph("")

    sample = [
        ("S001", "Alice Sharma", "2025CS001", "CSE", "5", "Mathematics, Physics", "alice@example.com"),
        ("S002", "Rahul Gupta", "2025CS002", "CSE", "5", "Mathematics, Chemistry", "rahul@example.com"),
        ("S003", "Priya Singh", "2025EC001", "ECE", "5", "Physics, Electronics", "priya@example.com"),
        ("S004", "Mohan Kumar", "2025ME001", "ME", "5", "Mathematics, Mechanics", "mohan@example.com"),
        ("S005", "Anjali Jain", "2025CE001", "CE", "5", "Physics, Environmental", "anjali@example.com"),
    ]

    for student_id, name, roll, department, sem, subjects, email in sample:
        doc.add_paragraph(f"Student ID: {student_id}")
        doc.add_paragraph(f"Name: {name}")
        doc.add_paragraph(f"Roll No: {roll}")
        doc.add_paragraph(f"Department: {department}")
        doc.add_paragraph(f"Semester: {sem}")
        doc.add_paragraph(f"Subjects: {subjects}")
        doc.add_paragraph(f"Email: {email}")
        doc.add_paragraph("")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_hostel_template_docx() -> bytes:
    """
    Generate a ready-to-use .docx template for hostel room upload.
    """
    from docx import Document as _Doc
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = _Doc()

    title = doc.add_heading("Hostel Room Upload Template", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "Instructions:\n"
        "• Each room entry must include Room ID, Floor, and Beds.\n"
        "• Separate each room entry with ONE blank line.\n"
        "• Floor: floor number (0 = ground floor, 1 = first floor, etc.)\n"
        "• Beds: number of beds in the room\n"
        "• Features are comma-separated (optional: AC, WiFi, Attached Bath, etc.)\n"
        "• Do NOT use tables — keep plain Key: Value format."
    )
    doc.add_paragraph("")

    # Sample entries
    entries = [
        ("101", "1", "2", "AC, WiFi, Attached Bath"),
        ("102", "1", "2", "AC, WiFi"),
        ("103", "1", "3", "AC, Shared Bath"),
        ("201", "2", "2", "AC, WiFi, Attached Bath"),
        ("202", "2", "2", "Non-AC, Shared Bath"),
        ("203", "2", "4", "AC, Shared Bath"),
        ("301", "3", "2", "AC, WiFi, Attached Bath"),
        ("302", "3", "3", "AC, WiFi, Shared Bath"),
    ]

    for room_id, floor, beds, features in entries:
        doc.add_paragraph(f"Room ID: {room_id}")
        doc.add_paragraph(f"Floor: {floor}")
        doc.add_paragraph(f"Beds: {beds}")
        doc.add_paragraph(f"Features: {features}")
        doc.add_paragraph("")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─── Hospital blueprint parser ────────────────────────────────────────────────

HOSPITAL_TYPE_ALIASES: dict[str, str] = {
    "ward": "ward",
    "general ward": "ward",
    "icu": "icu",
    "intensive care": "icu",
    "intensive care unit": "icu",
    "ot": "ot",
    "operation theatre": "ot",
    "operation theater": "ot",
    "operating theatre": "ot",
    "operating theater": "ot",
    "consultation": "consultation",
    "consult": "consultation",
    "consulting": "consultation",
    "consulting room": "consultation",
    "bed": "bed",
}

HOSPITAL_FIELD_ALIASES: dict[str, str] = {
    "room id": "room_id",
    "roomid": "room_id",
    "room no": "room_id",
    "room number": "room_id",
    "room": "room_id",
    "id": "room_id",
    "resource id": "room_id",
    "type": "type",
    "resource type": "type",
    "department": "department",
    "dept": "department",
    "specialty": "department",
    "speciality": "department",
    "capacity": "capacity",
    "beds": "capacity",
    "cap": "capacity",
    "floor": "floor",
    "floor number": "floor",
    "floor no": "floor",
    "facilities": "facilities",
    "equipment": "facilities",
    "amenities": "facilities",
    "features": "facilities",
}

REQUIRED_HOSPITAL_FIELDS = {"room_id", "type", "capacity"}


def _normalise_hospital_type(raw: str) -> Optional[str]:
    return HOSPITAL_TYPE_ALIASES.get(raw.strip().lower())


def _normalise_hospital_field_key(raw: str) -> Optional[str]:
    return HOSPITAL_FIELD_ALIASES.get(raw.strip().lower())


def _parse_hospital_block(lines: list[str]) -> dict:
    raw: dict = {}
    for line in lines:
        if ":" not in line:
            continue
        key_part, _, val_part = line.partition(":")
        field = _normalise_hospital_field_key(key_part)
        if field:
            raw[field] = val_part.strip()
    return raw


def validate_and_build_hospital_resource(raw: dict, block_index: int) -> dict:
    missing = REQUIRED_HOSPITAL_FIELDS - raw.keys()
    if missing:
        raise ParseError(
            f"Entry {block_index + 1}: missing required field(s): "
            f"{', '.join(sorted(missing))}"
        )

    try:
        capacity = int(str(raw["capacity"]).strip())
        if capacity <= 0:
            raise ValueError
    except ValueError:
        raise ParseError(
            f"Entry {block_index + 1}: 'Capacity' must be a positive integer, "
            f"got '{raw['capacity']}'"
        )

    normalised_type = _normalise_hospital_type(raw["type"])
    if normalised_type is None:
        raise ParseError(
            f"Entry {block_index + 1}: unrecognised Type '{raw['type']}'. "
            f"Allowed: Ward, ICU, OT, Consultation, Bed"
        )

    floor_raw = raw.get("floor", "0").strip() or "0"
    try:
        floor = int(floor_raw)
        if floor < 0:
            raise ValueError
    except ValueError:
        raise ParseError(
            f"Entry {block_index + 1}: 'Floor' must be a non-negative integer, "
            f"got '{floor_raw}'"
        )

    return {
        "room_id": raw["room_id"].strip(),
        "resource_type": normalised_type,
        "capacity": capacity,
        "department": raw.get("department", "").strip() or None,
        "floor": floor,
        "facilities": _parse_facilities(raw.get("facilities", "")),
    }


def parse_docx_hospital_resources(file_bytes: bytes) -> dict:
    """Parse a hospital blueprint .docx into validated resource dicts."""
    lines = extract_all_lines(file_bytes)
    raw_preview = [l for l in lines if not _is_blank(l)][:30]
    blocks = split_into_blocks(lines)

    resources: list[dict] = []
    errors: list[str] = []

    for idx, block in enumerate(blocks):
        if not any(":" in line for line in block):
            continue
        raw = _parse_hospital_block(block)
        if not raw:
            continue
        try:
            resources.append(validate_and_build_hospital_resource(raw, idx))
        except ParseError as exc:
            errors.append(str(exc))
        except Exception as exc:
            errors.append(f"Entry {idx + 1}: unexpected error — {exc}")

    if not resources and not errors:
        errors.append(
            "No valid hospital resource entries found. "
            "Use 'Key: Value' lines with Room ID, Type, and Capacity."
        )

    return {
        "resources": resources,
        "errors": errors,
        "total_blocks": len(blocks),
        "raw_preview": raw_preview,
    }


def generate_hospital_template_docx() -> bytes:
    """Ready-to-use hospital blueprint .docx template."""
    from docx import Document as _Doc
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = _Doc()
    title = doc.add_heading("Hospital Blueprint Upload Template", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "Instructions:\n"
        "• Each resource entry must include Room ID, Type, and Capacity.\n"
        "• Separate each entry with ONE blank line.\n"
        "• Valid Type values: Ward, ICU, OT, Consultation, Bed\n"
        "• Department, Floor, and Facilities are optional.\n"
        "• Facilities are comma-separated (e.g. monitors, ventilator).\n"
        "• Do NOT use tables — keep plain Key: Value format."
    )
    doc.add_paragraph("")

    entries = [
        ("W-101", "Ward", "Cardiology", "2", "10", "monitors, oxygen, nurse_call"),
        ("W-102", "Ward", "Orthopaedics", "2", "12", "oxygen, nurse_call"),
        ("ICU-01", "ICU", "Critical Care", "3", "6", "ventilator, monitors, defibrillator"),
        ("OT-1", "OT", "Surgery", "1", "1", "anesthesia, lights, monitors"),
        ("C-201", "Consultation", "General Medicine", "1", "1", "desk, examination_table"),
    ]

    for room_id, rtype, dept, floor, cap, facilities in entries:
        doc.add_paragraph(f"Room ID: {room_id}")
        doc.add_paragraph(f"Type: {rtype}")
        doc.add_paragraph(f"Department: {dept}")
        doc.add_paragraph(f"Floor: {floor}")
        doc.add_paragraph(f"Capacity: {cap}")
        doc.add_paragraph(f"Facilities: {facilities}")
        doc.add_paragraph("")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_lodge_template_docx() -> bytes:
    """
    Generate a ready-to-use .docx template for lodge room upload.
    """
    from docx import Document as _Doc
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = _Doc()

    title = doc.add_heading("Lodge Room Upload Template", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "Instructions:\n"
        "• Each room entry must include Room ID, Floor, and Beds.\n"
        "• Separate each room entry with ONE blank line.\n"
        "• Floor: floor number (0 = ground floor, 1 = first floor, etc.)\n"
        "• Beds: number of beds (sharing capacity per room)\n"
        "• Features are comma-separated (optional: AC, WiFi, TV, Kitchen, etc.)\n"
        "• Do NOT use tables — keep plain Key: Value format."
    )
    doc.add_paragraph("")

    # Sample entries
    entries = [
        ("A101", "1", "1", "AC, WiFi, TV, Attached Bath"),
        ("A102", "1", "2", "AC, WiFi, Shared Kitchen"),
        ("A103", "1", "3", "Non-AC, WiFi, Shared Bath"),
        ("B201", "2", "1", "AC, WiFi, TV, Attached Bath"),
        ("B202", "2", "2", "AC, WiFi, Attached Bath"),
        ("B203", "2", "4", "AC, Shared Kitchen, Shared Bath"),
        ("C301", "3", "1", "AC, WiFi, TV, Attached Bath"),
        ("C302", "3", "2", "AC, WiFi, Shared Bath"),
    ]

    for room_id, floor, beds, features in entries:
        doc.add_paragraph(f"Room ID: {room_id}")
        doc.add_paragraph(f"Floor: {floor}")
        doc.add_paragraph(f"Beds: {beds}")
        doc.add_paragraph(f"Features: {features}")
        doc.add_paragraph("")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
